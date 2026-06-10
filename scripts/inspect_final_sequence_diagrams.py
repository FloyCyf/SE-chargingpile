from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


DOCX = Path("D:/NeatDownload/第2次作业 概要设计模版最终版.docx")


def text_of(el):
    return "".join(node.text or "" for node in el.iter(qn("w:t"))).strip()


def main():
    doc = Document(DOCX)
    body = list(doc.element.body)
    current_heading = ""
    current_sub = ""
    drawings = []
    for bi, el in enumerate(body):
        if el.tag != qn("w:p"):
            continue
        txt = text_of(el)
        if txt.startswith("3."):
            current_heading = txt
        if txt.startswith("对象设计：") or txt.startswith("3."):
            current_sub = txt
        if list(el.iter(qn("w:drawing"))):
            drawings.append((bi, current_heading, current_sub, txt))
    print("drawings", len(drawings))
    for item in drawings:
        print(item)
    print("images", sum(1 for r in doc.part.rels.values() if "image" in r.reltype))
    print("tables", len(doc.tables), "paragraphs", len(doc.paragraphs))


if __name__ == "__main__":
    main()
